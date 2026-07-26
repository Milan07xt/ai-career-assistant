import os
import pdfplumber
import docx
import logging

logger = logging.getLogger(__name__)

def extract_text_from_file(file_obj, filename):
    """
    Extracts text content from PDF, DOCX, or TXT file objects.
    """
    ext = os.path.splitext(filename.lower())[1]
    text = ""
    
    try:
        if ext == ".pdf":
            with pdfplumber.open(file_obj) as pdf:
                pages_text = []
                for idx, page in enumerate(pdf.pages):
                    page_content = page.extract_text()
                    if page_content:
                        pages_text.append(page_content)
                text = "\n".join(pages_text)
                
        elif ext in [".docx", ".doc"]:
            doc = docx.Document(file_obj)
            paragraphs = [p.text for p in doc.paragraphs if p.text]
            # also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    paragraphs.append(" ".join(row_text))
            text = "\n".join(paragraphs)
            
        elif ext == ".txt":
            # read as bytes and decode
            file_obj.seek(0)
            text = file_obj.read()
            if isinstance(text, bytes):
                try:
                    text = text.decode("utf-8")
                except UnicodeDecodeError:
                    text = text.decode("latin-1")
                    
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
            
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {e}")
        text = f"[Extraction Error: Could not parse text from file. details: {str(e)}]"
        
    return text.strip()
