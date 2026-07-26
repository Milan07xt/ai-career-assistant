from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from rest_framework import generics, permissions, status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser
import docx

from .models import Resume, CoverLetter
from .serializers import ResumeSerializer
from .utils import extract_text_from_file
from ai_services.services import get_resume_analysis, generate_cover_letter

# =====================================================================
# HTML TEMPLATE VIEWS
# =====================================================================

@login_required
def analyze_view(request):
    if request.method == "POST" and request.FILES.get("resume_file"):
        uploaded_file = request.FILES["resume_file"]
        
        # 1. Save file in temporary DB object first to extract
        resume = Resume(user=request.user, file=uploaded_file)
        resume.save()
        
        # 2. Extract text
        extracted_text = extract_text_from_file(resume.file, uploaded_file.name)
        resume.extracted_text = extracted_text
        
        # 3. Call central AI service
        analysis = get_resume_analysis(extracted_text)
        resume.ats_score = analysis.get("ats_score", 70)
        resume.analysis_json = analysis
        resume.save()

        # 4. Automatically update User Profile if empty or update current resume
        profile = request.user.profile
        profile.resume_file = resume.file
        # Merge skills if user has none
        if not profile.skills:
            # Try to grab some skills from the resume or suggest
            profile.skills = ", ".join(analysis.get("missing_skills", ["Python", "Django"])[:3])
        profile.save()

        messages.success(request, "Resume analyzed successfully!")
        return redirect("resume_analyzer:detail", pk=resume.pk)
        
    return render(request, "resume_analyzer/upload.html")

@login_required
def history_view(request):
    resumes = Resume.objects.filter(user=request.user)
    return render(request, "resume_analyzer/history.html", {"resumes": resumes})

@login_required
def detail_view(request, pk):
    resume = get_object_or_404(Resume, pk=pk, user=request.user)
    analysis = resume.analysis_json or {}
    
    # Calculate scores
    ats_score = resume.ats_score
    ats_dashoffset = int(377 * (100 - ats_score) / 100)
    grammar_score = analysis.get("grammar_score", 80)
    formatting_score = analysis.get("formatting_score", 80)
    keyword_match = analysis.get("keyword_match", 70)
    
    context = {
        "resume": resume,
        "analysis": analysis,
        "ats_score": ats_score,
        "ats_dashoffset": ats_dashoffset,
        "grammar_score": grammar_score,
        "formatting_score": formatting_score,
        "keyword_match": keyword_match,
    }
    return render(request, "resume_analyzer/detail.html", context)


# =====================================================================
# REST API VIEWS
# =====================================================================

class APIResumeAnalyzeView(APIView):
    permission_classes = [permissions.IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, *args, **kwargs):
        if not request.FILES.get("file"):
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)
            
        uploaded_file = request.FILES["file"]
        
        # Save & Extract
        resume = Resume(user=request.user, file=uploaded_file)
        resume.save()
        
        extracted_text = extract_text_from_file(resume.file, uploaded_file.name)
        resume.extracted_text = extracted_text
        
        # AI analysis
        analysis = get_resume_analysis(extracted_text)
        resume.ats_score = analysis.get("ats_score", 70)
        resume.analysis_json = analysis
        resume.save()

        # Update profile
        profile = request.user.profile
        profile.resume_file = resume.file
        profile.save()

        serializer = ResumeSerializer(resume)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

class APIResumeHistoryView(generics.ListAPIView):
    permission_classes = [permissions.IsAuthenticated]
    serializer_class = ResumeSerializer

    def get_queryset(self):
        return Resume.objects.filter(user=self.request.user)

# =====================================================================
# COVER LETTER VIEWS
# =====================================================================

@login_required
def generate_cover_letter_view(request):
    if request.method == "POST":
        company_name = request.POST.get("company_name")
        job_desc = request.POST.get("job_description")
        
        # Grab latest resume text for AI context
        latest_resume = Resume.objects.filter(user=request.user).first()
        resume_text = latest_resume.extracted_text if latest_resume else ""
        if not resume_text:
            resume_text = f"Candidate Profile:\nPreferred Role: {request.user.profile.preferred_role}\nSkills: {request.user.profile.skills}\nBio: {request.user.profile.bio}"

        # Generate cover letter via AI
        content = generate_cover_letter(resume_text, job_desc, company_name)
        
        # Save Cover Letter
        cl = CoverLetter.objects.create(
            user=request.user,
            company_name=company_name,
            job_description=job_desc,
            content=content
        )
        messages.success(request, "Cover letter generated successfully!")
        return redirect("resume_analyzer:cover_letter_detail", cl_id=cl.id)
        
    return render(request, "resume_analyzer/cover_letter_form.html")

@login_required
def cover_letter_detail_view(request, cl_id):
    cl = get_object_or_404(CoverLetter, id=cl_id, user=request.user)
    return render(request, "resume_analyzer/cover_letter_detail.html", {"cl": cl})

@login_required
def download_docx_view(request, cl_id):
    cl = get_object_or_404(CoverLetter, id=cl_id, user=request.user)
    
    doc = docx.Document()
    # Add title and text
    doc.add_heading(f"Cover Letter - {cl.company_name}", level=1)
    doc.add_paragraph(f"Date: {cl.created_at.strftime('%B %d, %Y')}\n")
    
    # Body
    for para in cl.content.split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
            
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f'attachment; filename="cover_letter_{cl.company_name.replace(" ", "_")}.docx"'
    doc.save(response)
    return response

@login_required
def download_txt_view(request, cl_id):
    cl = get_object_or_404(CoverLetter, id=cl_id, user=request.user)
    
    response = HttpResponse(cl.content, content_type="text/plain")
    response["Content-Disposition"] = f'attachment; filename="cover_letter_{cl.company_name.replace(" ", "_")}.txt"'
    return response

# REST API Endpoint for Cover Letter
class APICoverLetterGenerateView(APIView):
    permission_classes = [permissions.IsAuthenticated]

    def post(self, request, *args, **kwargs):
        company_name = request.data.get("company_name")
        job_desc = request.data.get("job_description")
        
        if not company_name or not job_desc:
            return Response({"error": "company_name and job_description are required"}, status=status.HTTP_400_BAD_REQUEST)

        latest_resume = Resume.objects.filter(user=request.user).first()
        resume_text = latest_resume.extracted_text if latest_resume else ""
        if not resume_text:
            resume_text = f"Skills: {request.user.profile.skills}\nRole: {request.user.profile.preferred_role}"

        content = generate_cover_letter(resume_text, job_desc, company_name)
        
        cl = CoverLetter.objects.create(
            user=request.user,
            company_name=company_name,
            job_description=job_desc,
            content=content
        )
        
        return Response({
            "id": cl.id,
            "company_name": cl.company_name,
            "content": cl.content,
            "created_at": cl.created_at
        }, status=status.HTTP_201_CREATED)

